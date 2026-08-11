# -*- coding: utf-8 -*-
"""Laporan bagian 2: sequence diagram (gambar) + penjelasan langkah. Tanpa kode."""
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
DIAG = ROOT / "docs" / "diagrams"
OUT = ROOT / "docs" / "Laporan_Bagian_2_Flow_Sistem.docx"


def font(run, size=11, bold=False, color=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(15, 23, 42)
        run.font.name = "Calibri"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")


def para(doc, text, size=11, bold=False, justify=True, after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    font(r, size=size, bold=bold)
    return p


def bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(it, style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        for run in p.runs:
            run.font.size = Pt(11)
            run.font.name = "Calibri"


def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.space_before = Pt(4)
    r = p.add_run(text)
    font(r, size=10, color=(100, 116, 139))
    r.italic = True


def picture(doc, stem, width=6.4):
    path = DIAG / f"{stem}.png"
    if not path.exists():
        para(doc, f"(Gambar {stem} belum dirender.)", size=10, justify=False)
        return
    doc.add_picture(str(path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def flow_block(doc, title, stem, cap, steps):
    heading(doc, title, 2)
    picture(doc, stem)
    caption(doc, cap)
    para(doc, "Penjelasan singkat dari setiap langkah dalam flow:", bold=True, justify=False, after=6)
    bullets(doc, steps)
    doc.add_paragraph()


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.2)
    sec.right_margin = Cm(2.2)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("LAPORAN PROYEK — BAGIAN 2")
    font(r, size=12, bold=True, color=(15, 118, 110))

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Flow Sistem")
    font(r, size=26, bold=True, color=(15, 23, 42))

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Tolongin — sequence diagram dan penjelasan langkah")
    font(r, size=12, color=(71, 85, 105))

    heading(doc, "2. Flow Sistem", 1)
    para(
        doc,
        "Bagian ini menampilkan bagaimana pengguna, aplikasi React, server Express, "
        "basis data MySQL, dan gerbang pembayaran saling berinteraksi. Setiap flow "
        "memakai satu sequence diagram utuh, lalu penjelasan singkat per langkah. "
        "Pengguna tidak menulis ke basis data langsung: semua lewat antarmuka, "
        "lalu server, lalu penyimpanan. Sesi ditandai cookie aman. Uang pesanan "
        "ditahan sampai hasil kerja disetujui.",
    )

    flow_block(
        doc,
        "2.1 Alur autentikasi",
        "flow-auth",
        "Gambar 1. Sequence autentikasi: daftar, masuk, sesi, dan keluar.",
        [
            "Pengguna mengisi nama, email, HP, dan password di form daftar.",
            "Aplikasi memastikan konfirmasi password cocok sebelum dikirim.",
            "Server memeriksa email belum terpakai, menyimpan akun, dan mengamankan password.",
            "Server memasang cookie sesi. Pengguna diarahkan melengkapi profil.",
            "Saat masuk, server mencocokkan password. Benar: cookie baru dan masuk beranda. Salah: peringatan.",
            "Halaman privat menanyakan sesi ke server. Token sah: halaman tampil. Tidak sah: kembali ke login.",
            "Keluar menghapus cookie di server, lalu pengguna kembali ke beranda.",
        ],
    )

    flow_block(
        doc,
        "2.2 Alur data",
        "flow-data",
        "Gambar 2. Sequence alur data: permintaan berhasil atau ditolak.",
        [
            "Pengguna mengklik atau mengirim form di antarmuka.",
            "Aplikasi mengirim permintaan ke server beserta cookie sesi.",
            "Server memeriksa login dan hak akses (termasuk verifikasi jika diperlukan).",
            "Jika belum masuk, pengguna diarahkan login. Jika syarat kurang, diarahkan verifikasi atau ditolak.",
            "Jika isian salah, muncul peringatan validasi. Jika data tidak ada, muncul tidak ditemukan.",
            "Jika lolos, server membaca atau menulis ke basis data, lalu mengirim hasil.",
            "Antarmuka menampilkan daftar, kartu, atau memindahkan pengguna ke halaman berikutnya.",
        ],
    )

    flow_block(
        doc,
        "2.3 Interaksi utama pengguna: sewa jasa",
        "flow-sewa",
        "Gambar 3. Sequence sewa jasa: permintaan, bayar, pengerjaan, sampai dana cair atau sengketa.",
        [
            "Pembeli menyewa jasa yang sedang tayang. Server membuat pesanan menunggu konfirmasi penjual.",
            "Penjual mendapat pemberitahuan, lalu menolak atau menerima.",
            "Jika ditolak, alur berhenti. Jika diterima, pembeli membayar lewat gerbang pembayaran.",
            "Setelah lunas, uang ditahan dan status menjadi sedang dikerjakan. Ruang kerja terbuka.",
            "Penjual mengunggah hasil. Pembeli dapat minta revisi, menyetujui, atau membuka sengketa.",
            "Setuju: pesanan selesai dan dana cair ke penjual. Sengketa: admin yang memutuskan.",
            "Selama pesanan belum selesai, pemilik jasa tidak dapat menonaktifkan atau menghapus listing itu.",
        ],
    )

    flow_block(
        doc,
        "2.4 Interaksi utama pengguna: lowongan dan lamaran",
        "flow-kerja",
        "Gambar 4. Sequence lowongan: pasang, lamar, terima satu orang, atau tutup rekrutmen.",
        [
            "Klien memasang lowongan. Lowongan terbuka dan muncul di Cari Kerja.",
            "Pelamar mengirim penawaran. Lamaran berstatus menunggu. Klien mendapat pemberitahuan.",
            "Setelah ada pelamar, isi lowongan tidak dapat diubah agar penawaran tetap adil.",
            "Jika klien menerima satu orang: lowongan terisi, pelamar lain otomatis tidak terpilih, lalu dibuat pesanan. Pembayaran mengikuti alur sewa jasa.",
            "Jika klien menutup lowongan: semua yang masih menunggu ditolak, lowongan batal.",
            "Hapus hanya boleh jika belum ada pelamar. Jika sudah ada pekerja atau proyek berjalan, listing terkunci sampai pesanan selesai.",
        ],
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("— akhir bagian 2 —")
    font(r, size=10, color=(148, 163, 184))

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
