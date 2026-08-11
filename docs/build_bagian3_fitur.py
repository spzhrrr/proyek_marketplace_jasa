# -*- coding: utf-8 -*-
"""Laporan bagian 3: gambaran fitur. Bahasa manusia, istilah Inggris tetap."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Laporan_Bagian_3_Gambaran_Fitur.docx"


def font(run, size=11, bold=False, color=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(15, 23, 42)
        run.font.name = "Calibri"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")


def para(doc, text, size=11, bold=False, justify=True, after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    font(r, size=size, bold=bold)
    return p


def bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(it, style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        for run in p.runs:
            run.font.size = Pt(11)
            run.font.name = "Calibri"


def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            if p.runs:
                p.runs[0].bold = True
                p.runs[0].font.size = Pt(10)
                p.runs[0].font.name = "Calibri"
                p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
        shade_cell(cell, "0F172A")
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = t.rows[r_i + 1].cells[c_i]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = "Calibri"
            if r_i % 2 == 1:
                shade_cell(cell, "F8FAFC")
    doc.add_paragraph()


def fitur(doc, title, body):
    para(doc, title, bold=True, justify=False, after=4)
    para(doc, body)


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.2)
    sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.3)
    sec.right_margin = Cm(2.3)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("LAPORAN PROYEK — BAGIAN 3")
    font(r, size=12, bold=True, color=(15, 118, 110))

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Gambaran Keseluruhan Fitur")
    font(r, size=24, bold=True, color=(15, 23, 42))

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Tolongin — marketplace jasa dan lowongan")
    font(r, size=12, color=(71, 85, 105))

    heading(doc, "3. Gambaran Keseluruhan Fitur", 1)
    para(
        doc,
        "Tolongin itu dual marketplace. Satu sisi orang jual jasa, sisi lain orang buka "
        "lowongan lalu pilih pelamar. Fitur yang dibangun mengikuti kebutuhan itu: akun, "
        "katalog, transaksi, chat, dan panel admin. Yang ditulis di sini fitur yang memang "
        "ada di aplikasi, bukan wishlist.",
    )

    heading(doc, "3.1 Fitur Utama yang Dikembangkan", 2)
    para(
        doc,
        "Intinya aplikasi harus bisa dipakai end-to-end: daftar, masuk, pasang listing, "
        "ada yang sewa atau lamar, bayar, kerja, baru dana cair. Tanpa salah satu dari "
        "itu, marketplace-nya terasa setengah jadi. Daftar di bawah ini yang dianggap fitur inti.",
    )
    table(
        doc,
        ["Fitur inti", "Yang dikerjakan di Tolongin"],
        [
            ["Register & login", "Form daftar/masuk, password di-hash, session lewat cookie JWT httpOnly"],
            ["Manajemen sesi", "GET /me, logout hapus cookie, halaman privat lewat ProtectedRoute"],
            ["CRUD jasa", "Buat, baca, edit, nonaktif, hapus (dengan constraint kalau masih ada sewa)"],
            ["CRUD lowongan", "Pasang, edit (sebelum ada pelamar), tutup, hapus, tampil/sembunyi di katalog"],
            ["Katalog publik", "Cari Jasa dan Cari Kerja: list, filter kategori, detail"],
            ["Sewa jasa", "Request PENDING, penjual terima/tolak, lanjut bayar"],
            ["Lamaran kerja", "Apply, terima satu orang, sisanya auto-reject, atau Tutup lowongan"],
            ["Escrow & payment", "Bayar lewat gateway mock, uang ditahan, cair setelah hasil disetujui"],
            ["Verifikasi identitas", "OTP email/HP, review KTP dan rekening oleh admin"],
            ["Dashboard user", "Listing sendiri, pesanan, lamaran, saldo wallet"],
            ["Chat & notifikasi", "Thread per listing, lonceng di header"],
            ["Profil & portfolio", "Storefront publik plus file karya"],
            ["Admin panel", "Antrian KTP/bank, user, pesanan, withdrawal, laporan"],
        ],
    )
    para(
        doc,
        "Selain itu ada onboarding lengkapi profil, review setelah order selesai, dan "
        "withdrawal saldo. Bukan fitur 'wajib silabus' satu-satu, tapi tanpa itu alur "
        "jual-beli terasa putus.",
    )

    heading(doc, "3.2 Deskripsi Fitur", 2)
    para(
        doc,
        "Tiap fitur dijelaskan dua hal: cara kerjanya di aplikasi, dan kenapa itu mirip "
        "dengan yang dipakai di kerja nyata (freelance platform, job board, atau internal tool).",
    )

    fitur(
        doc,
        "Register, login, dan session",
        "User daftar dengan nama, email, HP, dan password. Konfirmasi password dicek di frontend "
        "dulu supaya tidak bolak-balik ke server cuma karena typo. Password disimpan hash bcrypt, "
        "bukan plain text. Setelah register atau login, server pasang cookie JWT yang httpOnly — "
        "JavaScript tidak bisa baca token itu, jadi lebih aman dari XSS biasa dibanding nyimpan "
        "token di localStorage. Halaman yang butuh login memanggil session lewat GET /me. Logout "
        "menghapus cookie. Di lingkungan kerja, pola ini dekat dengan web app kantoran: login sekali, "
        "session ikut di browser, bukan API key yang di-paste user.",
    )
    fitur(
        doc,
        "Onboarding profil",
        "Habis daftar, user disuruh lengkapi foto, bio, dan kota. Tanpa ini, kartu katalog dan "
        "profil publik cuma nama kosong. Di kerja nyata hampir semua product onboarding seperti ini: "
        "akun sudah ada, tapi belum 'siap dipakai' sampai profil minimal terisi.",
    )
    fitur(
        doc,
        "Verifikasi bertingkat (email, HP, KTP, bank)",
        "Bukan langsung bebas transaksi. Email dan HP lewat OTP. KTP di-upload, admin yang approve "
        "atau tolak. Rekening juga: nama pemegang harus sama dengan KTP, lalu admin review. Sewa "
        "dan lamar butuh KTP approved. Post jasa dan tarik dana butuh rekening approved. Ini versi "
        "ringan dari KYC di marketplace sungguhan. Di kantor, mirip maker-checker: user submit, "
        "orang ops yang mengizinkan.",
    )
    fitur(
        doc,
        "Katalog jasa dan lowongan",
        "Cari Jasa (identitas biru) dan Cari Kerja (ungu) terpisah supaya orang tidak nyasar. "
        "Listing publik hanya yang aktif. Ada filter kategori fisik/digital dan search. Kalau "
        "database kosong, yang tampil empty state, bukan error. Ini discovery layer — sama seperti "
        "orang buka job board atau Fiverr: lihat dulu, baru klik detail.",
    )
    fitur(
        doc,
        "CRUD listing jasa",
        "Penjual bisa pasang jasa (judul, kategori, harga, estimasi hari, cover, skill), edit, "
        "sembunyikan dari katalog, atau hapus. Bedanya dengan CRUD kampus yang bebas delete: "
        "kalau masih ada permintaan sewa atau pesanan berjalan, edit / nonaktif / hapus dikunci. "
        "Kalau sudah ada riwayat pesanan selesai, hapus jadi archive, bukan hard delete, supaya "
        "order lama tidak yatim. Di kerja nyata, product listing hampir tidak pernah boleh dihapus "
        "sembarangan kalau sudah ada invoice.",
    )
    fitur(
        doc,
        "CRUD listing lowongan",
        "Klien pasang lowongan: budget, deadline lamaran, tag urgent kalau butuh hari ini. "
        "Edit hanya selama belum ada pelamar — supaya orang yang sudah apply tidak kena ubah harga "
        "di tengah jalan. Nonaktifkan katalog beda dengan Tutup. Tutup = tolak semua yang masih "
        "PENDING. Hapus hanya jika belum ada pelamar sama sekali. Yang sudah di-hire diselesaikan "
        "lewat order, bukan dengan menghapus lowongan. Pola ini dekat dengan ATS (applicant tracking) "
        "kecil: job posting punya status, bukan cuma tombol hapus.",
    )
    fitur(
        doc,
        "Sewa jasa",
        "Pembeli kirim request. Penjual terima atau tolak. Baru setelah diterima, pembeli bayar. "
        "Tidak langsung checkout seperti e-commerce barang. Masuk akal untuk jasa: penjual masih "
        "bisa bilang tidak sanggup. Di kerja freelance, ini tahap accept request sebelum invoice.",
    )
    fitur(
        doc,
        "Lamaran kerja",
        "Pelamar kirim cover letter, harga (batas 50–150% budget), dan file. Klien cuma boleh "
        "terima satu orang. Saat diterima, pelamar lain otomatis tidak terpilih. Klien juga bisa "
        "tolak manual dengan alasan, atau Tutup lowongan. Relevansinya jelas: hiring tidak boleh "
        "dobel worker untuk satu job yang sama, dan kandidat yang kalah harus dapat status, bukan "
        "digantung.",
    )
    fitur(
        doc,
        "Escrow, payment mock, revisi, sengketa",
        "Uang tidak langsung ke penjual. Setelah bayar, webhook menandai lunas lalu escrow HELD "
        "dan pekerjaan mulai. Penjual unggah bukti. Pembeli boleh minta revisi (maks tiga), setujui, "
        "atau buka sengketa ke admin. Setuju: COMPLETED, dana cair ke wallet. Ini meniru escrow "
        "di marketplace jasa. Gateway-nya masih mock (bukan Midtrans), tapi alur statusnya sudah "
        "dipisah: UNPAID, HELD, RELEASED. Di kerja, orang finance tidak campur 'sudah bayar' dengan "
        "'sudah boleh dicairkan'.",
    )
    fitur(
        doc,
        "Dashboard, chat, notifikasi",
        "Dashboard adalah beranda setelah login: jasa/lowongan sendiri, pesanan, lamaran, saldo. "
        "Chat menempel di listing (bukan chat global campur semua orang). Notifikasi muncul di "
        "lonceng kalau ada sewa, lamaran, terima, bayar, atau lowongan ditutup. Di kantor ini "
        "mirip inbox + activity feed: orang tidak harus refresh halaman terus untuk tahu ada kerjaan baru.",
    )
    fitur(
        doc,
        "Profil publik, portfolio, review",
        "Profil orang lain menampilkan jasa/lowongan yang tayang, portfolio, ulasan, dan kerjaan "
        "selesai. Portfolio pakai kategori. Review muncul setelah order selesai. Ini storefront "
        "freelancer: orang putuskan sewa atau lamar setelah lihat jejak kerja, bukan cuma harga.",
    )
    fitur(
        doc,
        "Admin panel dan withdrawal",
        "Admin bukan user biasa. Mereka antri KTP/bank, lihat pesanan, laporan user, dan permintaan "
        "tarik saldo. Withdrawal cek rekening sudah verified dan saldo cukup. Di kerja nyata ini "
        "ops / trust & safety: ada manusia yang pegang kasus yang tidak bisa diputus mesin.",
    )

    heading(doc, "3.3 Skalabilitas dan Pengembangan Lanjutan", 2)
    para(
        doc,
        "Sekarang aplikasi jalan di localhost, MySQL satu database, payment masih mock, OTP "
        "belum lewat SMS/email provider sungguhan. Cukup untuk tugas dan demo alur. Kalau mau "
        "dipakai lebih serius, yang paling kelihatan dulu bukan tambah fitur aneh, tapi "
        "mengganti bagian yang masih tiruan dan yang akan nge-hang kalau usernya nambah.",
    )
    para(doc, "Yang masuk akal dikerjakan berikutnya:", bold=True, justify=False, after=4)
    bullets(
        doc,
        [
            "Payment gateway asli (Midtrans atau Xendit) menggantikan mock, plus webhook yang sama polanya.",
            "OTP lewat email SMTP / SMS provider, bukan kode yang cuma tampil di UI development.",
            "Deploy ke cloud (app + MySQL managed). Upload file ke object storage, jangan numpuk di disk server.",
            "Index dan pagination katalog yang lebih ketat; search judul bisa ditambah full-text kalau data ramai.",
            "Rate limit di login dan OTP supaya tidak gampang di-spam.",
            "Cache ringan (Redis) untuk session atau hitungan dashboard, kalau query mulai berat.",
            "Notifikasi realtime (WebSocket) supaya chat tidak perlu refresh.",
            "Unit test dan integration test untuk order/escrow — bagian ini paling gampang rusak kalau diubah.",
            "PWA atau mobile view yang lebih rapi; sekarang sudah web, tapi belum jadi app store.",
        ],
    )
    para(
        doc,
        "Yang sengaja tidak dipaksa sekarang: machine learning rekomendasi, chat video, atau "
        "multi-currency. Belum ada datanya, dan alur escrow saja sudah cukup rawan. Lebih berguna "
        "mengunci status pesanan dan constraint listing supaya tidak kacau, baru scale ke luar.",
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("— akhir bagian 3 —")
    font(r, size=10, color=(148, 163, 184))

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
