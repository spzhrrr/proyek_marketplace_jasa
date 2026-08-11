# -*- coding: utf-8 -*-
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Laporan_Proyek_Tolongin.docx"
OUT_FALLBACK = ROOT / "docs" / "Laporan_Proyek_Tolongin_baru.docx"
OUT_FALLBACK2 = ROOT / "docs" / "Laporan_Proyek_Tolongin_flow.docx"
SHOT = ROOT / "docs" / "screenshots"
DIAG = ROOT / "docs" / "diagrams"


def set_run_font(run, name="Calibri", size=11, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(15, 23, 42)
        run.font.name = "Calibri"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    return p


def para(doc, text, size=11, bold=False, space_after=8, justify=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.space_before = Pt(2)
    run = p.add_run(text)
    set_run_font(run, size=9, color=(100, 116, 139))
    run.italic = True


def add_picture(doc, path, width=6.2):
    path = Path(path)
    if not path.exists():
        para(doc, f"(file belum ada: {path.name})", size=9, justify=False)
        return
    doc.add_picture(str(path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def shot(doc, name, cap, width=6.2):
    add_picture(doc, SHOT / f"{name}.png", width)
    caption(doc, cap)


def mermaid_block(doc, filename):
    src = DIAG / filename
    text = src.read_text(encoding="utf-8") if src.exists() else ""
    para(
        doc,
        f"Kode Mermaid ({filename}) — ubah di docs/diagrams lalu render ulang, atau tempel ke mermaid.live:",
        size=9,
        bold=True,
        justify=False,
        space_after=4,
    )
    t = doc.add_table(rows=1, cols=1)
    cell = t.rows[0].cells[0]
    shade_cell(cell, "F1F5F9")
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text.strip())
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(15, 23, 42)
    doc.add_paragraph()


def diagram(doc, stem, cap):
    add_picture(doc, DIAG / f"{stem}.png", 6.0)
    caption(doc, cap)
    mermaid_block(doc, f"{stem}.mmd")


def flow_points(doc, diagrams, explanations):
    para(doc, "Poin 1 — Sequence diagram (aksi utama)", bold=True, justify=False, space_after=4)
    bullet(doc, diagrams)
    para(doc, "Poin 2 — Penjelasan", bold=True, justify=False, space_after=4)
    bullet(doc, explanations)


def diagram_with_steps(doc, stem, cap, steps):
    diagram(doc, stem, cap)
    para(doc, "Penjelasan singkat dari setiap langkah:", bold=True, justify=False, space_after=4)
    bullet(doc, steps)


def table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            if p.runs:
                p.runs[0].bold = True
                p.runs[0].font.size = Pt(9)
                p.runs[0].font.name = "Calibri"
                p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
        shade_cell(cell, "0F172A")
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = t.rows[r_i + 1].cells[c_i]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.name = "Calibri"
            if r_i % 2 == 1:
                shade_cell(cell, "F8FAFC")
    doc.add_paragraph()


def bullet(doc, items):
    for it in items:
        p = doc.add_paragraph(it, style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = Pt(11)
            run.font.name = "Calibri"


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.2)
    sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.3)
    sec.right_margin = Cm(2.3)

    for _ in range(2):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("LAPORAN PROYEK AKHIR")
    set_run_font(r, size=13, bold=True, color=(15, 118, 110))

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Tolongin")
    set_run_font(r, size=28, bold=True, color=(15, 23, 42))

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Marketplace Jasa & Lowongan Kerja")
    set_run_font(r, size=14, color=(71, 85, 105))

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Kategori Aplikasi — Fullstack JavaScript")
    set_run_font(r, size=12)

    doc.add_paragraph()
    for line in [
        "Kerangka mengikuti Exam 3 (9 bagian laporan proyek).",
        "Stack: Express.js · React.js (Vite) · MySQL (mysql2)",
        "Repo: https://github.com/spzhrrr/proyek_marketplace_jasa",
        "Lokal: http://localhost:5173  |  API http://localhost:3000",
        "Screenshot diambil dari web app yang berjalan (Agustus 2026).",
        "Diagram dari kode Mermaid di folder docs/diagrams/ (bukan gambar AI).",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        set_run_font(r, size=11, color=(51, 65, 85))

    doc.add_page_break()

    add_heading(doc, "1. Pendahuluan Proyek", 1)
    add_heading(doc, "1.1 Deskripsi singkat", 2)
    para(
        doc,
        "Tolongin adalah marketplace dua sisi. Satu sisi untuk orang yang jual jasa "
        "(desain, coding, servis AC, bersih-bersih, dan sejenisnya). Sisi lain untuk orang "
        "yang buka lowongan kerja lepas, lalu pilih pelamar. Dua-duanya butuh pembayaran "
        "yang tidak langsung cair ke penjual — uang ditahan (escrow) sampai hasil kerja disetujui.",
    )
    para(
        doc,
        "Masalah yang mau diatasi sederhana: di grup chat atau postingan acak, harga tidak jelas, "
        "riwayat orangnya susah dicek, dan tipu-menipu pembayaran masih sering. Aplikasi ini "
        "mengikat identitas (email, HP, KTP), katalog, chat per listing, dan alur pesanan dalam satu tempat.",
    )
    para(
        doc,
        "Penggunanya tiga. User biasa (bisa jadi pembeli, penjual, atau pelamar). Admin yang "
        "menyetujui KTP/rekening, lihat sengketa, dan kelola laporan. Sistem gateway mock untuk "
        "mensimulasikan pembayaran tanpa Midtrans sungguhan.",
    )
    shot(doc, "01-beranda", "Gambar 1. Beranda Tolongin (screenshot nyata, localhost:5173).")

    add_heading(doc, "1.2 Lingkup aplikasi", 2)
    para(doc, "Fitur inti (wajib ada supaya aplikasi terasa lengkap):", bold=True, justify=False)
    bullet(
        doc,
        [
            "Auth: register, login, logout; password bcrypt; sesi JWT di cookie httpOnly.",
            "CRUD jasa dan lowongan + katalog publik (filter, detail).",
            "Transaksi: sewa jasa, lamaran kerja, escrow, revisi, review.",
            "Verifikasi bertingkat: email OTP, HP OTP, KTP (admin), rekening (admin).",
            "Dashboard user, chat, notifikasi, profil publik + portfolio.",
            "Panel admin: antrian KTP/bank, pengguna, pesanan, withdrawal, laporan.",
        ],
    )
    para(doc, "Spesifikasi teknis yang dipakai:", bold=True, justify=False)
    table(
        doc,
        ["Bagian", "Pilihan"],
        [
            ["Frontend", "React 18 + Vite 6, React Router, CSS sendiri (tanpa Redux)"],
            ["Backend", "Express 4, MVC (routes / controllers / models)"],
            ["Database", "MySQL, skema proyek_marketplace, driver mysql2 (bukan Sequelize)"],
            ["Auth", "bcryptjs + jsonwebtoken, cookie httpOnly"],
            ["Upload", "Multer (cover jasa, KTP, bukti kerja, portfolio)"],
            ["Pembayaran", "Mock payment gateway + webhook PAID"],
        ],
    )
    para(
        doc,
        "Catatan silabus: PostgreSQL boleh diganti MySQL. EJS dan Redux tidak dipakai; UI React "
        "dan state auth lewat Context. Deploy cloud belum — masih lokal.",
    )

    add_heading(doc, "2. Flow Sistem", 1)
    add_heading(doc, "2.1 Ringkasan komponen", 2)
    para(
        doc,
        "Pengguna dan admin hanya lewat React. Aksi data dikirim JSON ke Express. Cookie JWT "
        "httpOnly menandai sesi. Model mysql2 menulis ke MySQL. Gateway mock dipakai saat bayar; "
        "hasilnya kembali ke API lewat webhook PAID.",
    )
    diagram(doc, "alur-sistem", "Gambar 2. Interaksi komponen user, frontend, backend, database, gateway.")

    add_heading(doc, "2.2 Alur autentikasi", 2)
    flow_points(
        doc,
        [
            "Sequence register — daftar akun baru",
            "Sequence login — masuk dengan email dan password",
            "Sequence sesi & logout — GET /me dan hapus cookie",
        ],
        [
            "Password di-hash bcrypt; JWT tidak disimpan di localStorage.",
            "Login gagal mengembalikan 401 “Email atau password salah”.",
            "Halaman privat selalu cek sesi lewat cookie httpOnly.",
            "Logout menghapus cookie di server.",
        ],
    )
    diagram_with_steps(
        doc,
        "auth-register",
        "Gambar 3. Sequence register.",
        [
            "User isi nama, email, HP, password di form React.",
            "Frontend cek konfirmasi password sebelum kirim.",
            "POST /api/mp/auth/register: email unik, bcrypt.hash, INSERT users.",
            "Server Set-Cookie token JWT httpOnly.",
            "React arahkan onboarding (lengkapi profil).",
        ],
    )
    diagram_with_steps(
        doc,
        "auth-login",
        "Gambar 4. Sequence login.",
        [
            "User kirim email dan password.",
            "POST /api/mp/auth/login: ambil user, bcrypt.compare.",
            "Sukses: Set-Cookie JWT, masuk dashboard/beranda.",
            "Gagal: 401, UI menampilkan Alert.",
        ],
    )
    diagram_with_steps(
        doc,
        "auth-sesi",
        "Gambar 5. Sequence sesi dan logout.",
        [
            "Halaman privat memanggil GET /api/mp/me dengan cookie.",
            "JWT valid: session user dikembalikan, halaman dirender.",
            "Token hilang/rusak: 401, redirect /login.",
            "Logout: POST /api/mp/auth/logout menghapus cookie.",
        ],
    )

    add_heading(doc, "2.3 Alur data", 2)
    flow_points(
        doc,
        [
            "Sequence request sukses — UI sampai MySQL lalu JSON",
            "Sequence request gagal — 401 / 403 / 400 / 404",
        ],
        [
            "Semua fetch lewat api.js dengan credentials include agar cookie ikut.",
            "SQL di model memakai parameterized query.",
            "Response seragam: { ok, data } atau { ok:false, error }.",
        ],
    )
    diagram_with_steps(
        doc,
        "data-sukses",
        "Gambar 6. Sequence alur data sukses.",
        [
            "User klik atau submit di UI.",
            "React fetch /api/mp/... + cookie.",
            "Guard JWT dan validasi lolos.",
            "Model menjalankan SQL ke MySQL.",
            "JSON ok:true; React render hasil.",
        ],
    )
    diagram_with_steps(
        doc,
        "data-gagal",
        "Gambar 7. Sequence alur data gagal.",
        [
            "401: belum login → Alert dan /login.",
            "403: KTP/bank/role kurang → halaman verifikasi.",
            "400: input salah → Alert validasi.",
            "404: resource tidak ada → pesan tidak ditemukan.",
        ],
    )

    add_heading(doc, "2.4 Interaksi utama: sewa jasa", 2)
    flow_points(
        doc,
        [
            "Sequence minta sewa — pembeli kirim request, penjual terima/tolak",
            "Sequence bayar — gateway mock dan webhook PAID",
            "Sequence selesai — bukti kerja, revisi, cair dana, atau sengketa",
        ],
        [
            "Uang tidak langsung ke penjual; ditahan escrow HELD sampai hasil disetujui.",
            "Revisi maksimal 3 kali.",
            "Selama pesanan hidup, jasa tidak boleh dinonaktifkan atau dihapus.",
        ],
    )
    diagram_with_steps(
        doc,
        "sewa-minta",
        "Gambar 8. Sequence permintaan sewa.",
        [
            "Pembeli klik Sewa pada jasa aktif.",
            "POST /api/mp/jasa/:id/sewa membuat orders PENDING.",
            "Penjual dapat notifikasi.",
            "Tolak → REJECTED. Terima → ACCEPTED, pembeli siap bayar.",
        ],
    )
    diagram_with_steps(
        doc,
        "sewa-bayar",
        "Gambar 9. Sequence pembayaran escrow.",
        [
            "Pembeli POST /orders/:id/bayar.",
            "API membuat transaksi di gateway mock.",
            "Pembeli simulasi pay.",
            "Webhook PAID: escrow HELD, status IN_PROGRESS, ruang kerja terbuka.",
        ],
    )
    diagram_with_steps(
        doc,
        "sewa-selesai",
        "Gambar 10. Sequence penyelesaian pesanan.",
        [
            "Penjual unggah bukti kerja.",
            "Pembeli minta revisi (maks 3), setujui, atau buka sengketa.",
            "Setujui → COMPLETED dan dana cair ke wallet.",
            "Sengketa → DISPUTED, admin yang putuskan.",
        ],
    )

    add_heading(doc, "2.5 Interaksi utama: lowongan dan lamaran", 2)
    flow_points(
        doc,
        [
            "Sequence pasang & lamar — klien buka lowongan, pelamar kirim penawaran",
            "Sequence terima pelamar — satu orang diterima, yang lain AUTO_FILLED",
            "Sequence tutup / hapus — batalkan rekrutmen sesuai constraint",
        ],
        [
            "Hanya satu pelamar yang boleh diterima per lowongan.",
            "Setelah ada pelamar, edit judul/budget dikunci.",
            "Hapus hanya jika belum ada pelamar; Tutup menolak semua PENDING.",
        ],
    )
    diagram_with_steps(
        doc,
        "kerja-lamar",
        "Gambar 11. Sequence pasang lowongan dan kirim lamaran.",
        [
            "Klien POST /api/mp/lowongan → jobs OPEN, tampil di Cari Kerja.",
            "Pelamar POST /lowongan/:id/lamar (cover letter, harga, file).",
            "Lamaran berstatus PENDING; klien dapat notifikasi.",
            "Setelah ada pelamar, edit lowongan dikunci.",
        ],
    )
    diagram_with_steps(
        doc,
        "kerja-terima",
        "Gambar 12. Sequence terima satu pelamar.",
        [
            "Klien POST /applications/:id/terima.",
            "Pelamar itu ACCEPTED; job FILLED.",
            "Pelamar lain otomatis AUTO_FILLED.",
            "Order escrow dibuat; lanjut bayar seperti sewa jasa.",
        ],
    )
    diagram_with_steps(
        doc,
        "kerja-tutup",
        "Gambar 13. Sequence tutup atau hapus lowongan.",
        [
            "Tutup: semua PENDING jadi JOB_CLOSED, job CANCELLED.",
            "Hapus: hanya jika belum ada pelamar.",
            "Sudah ada hire/pesanan hidup: API 400, harus diselesaikan dulu.",
        ],
    )

    add_heading(doc, "3. Gambaran Keseluruhan Fitur", 1)
    add_heading(doc, "3.1 Fitur utama", 2)
    table(
        doc,
        ["Fitur", "Cara kerja di aplikasi", "Kaitan kerja nyata"],
        [
            ["Login / register", "Form + validasi + cookie JWT", "Sesi user, bukan token di localStorage"],
            ["Lengkapi profil", "Foto, bio, kota; onboarding gate", "Identitas sebelum transaksi"],
            ["Verifikasi 4 tahap", "OTP email/HP, KTP & bank review admin", "KYC ringan marketplace"],
            ["Katalog jasa/kerja", "List, filter, kartu, empty state", "Discovery seperti job board"],
            ["CRUD listing", "Buat/edit/nonaktif; guard KTP/bank", "Seller dashboard"],
            ["Sewa & lamaran", "Request atau apply; satu hire per job", "Hiring + order"],
            ["Escrow + bayar mock", "HELD sampai hasil disetujui", "Escrow marketplace"],
            ["Chat & notifikasi", "Thread per listing + lonceng", "Komunikasi deal"],
            ["Profil & portfolio", "Kartu listing + file karya", "Storefront freelancer"],
            ["Admin", "Antrian KTP/bank, user, order, laporan", "Ops / trust & safety"],
        ],
    )

    add_heading(doc, "3.2 Deskripsi fitur (yang kelihatan di UI)", 2)
    para(
        doc,
        "Katalog jasa memakai identitas biru; katalog kerja ungu. Search dock di header tetap abu-abu "
        "netral. Beranda user (dashboard) memakai teal supaya tidak meniru warna jasa/kerja. "
        "Saat database masih kosong, katalog menampilkan empty state — bukan error.",
    )
    shot(doc, "06-katalog-jasa", "Gambar 14. Cari Jasa — empty state nyata (0 listing di DB).")
    shot(doc, "08-katalog-kerja", "Gambar 15. Cari Kerja — identitas ungu, 0 lowongan.")
    shot(doc, "11-dashboard", "Gambar 16. Beranda Saya: saldo, escrow, tab jasa/lowongan/pesanan.")
    shot(doc, "12-verifikasi", "Gambar 17. Hub verifikasi (email/HP/KTP selesai; rekening masih tahap 4).")
    shot(doc, "20-post-jasa", "Gambar 18. Form post jasa (judul, kategori, harga, cover).")
    shot(doc, "21-post-lowongan", "Gambar 19. Form post lowongan (budget, urgent, deadline lamaran).")
    shot(doc, "16-admin", "Gambar 20. Panel admin — antrian verifikasi dan ringkasan.")

    add_heading(doc, "3.3 Skalabilitas dan lanjutan", 2)
    para(
        doc,
        "Sekarang query masih langsung di model, cukup untuk tugas. Kalau traffic naik: indeks "
        "katalog (category, is_active, created_at), pagination yang sudah ada dipertahankan, "
        "upload dipindah ke object storage, dan gateway diganti Midtrans/Xendit. "
        "Fitur lanjutan yang masuk akal: ranking katalog dari rating, notifikasi email sungguhan, "
        "dan deploy (Railway/Vercel + MySQL managed).",
    )

    add_heading(doc, "4. Struktur Database", 1)
    add_heading(doc, "4.1 ERD", 2)
    para(
        doc,
        "Skema MySQL nama proyek_marketplace. Relasi utama 1–N. Satu user banyak jasa/lowongan/"
        "lamaran/order. Satu lowongan banyak lamaran. Satu order satu payment PAID (constraint unik). "
        "File override lengkap: backend/database/phpmyadmin_full_override.sql.",
    )
    diagram(doc, "erd", "Gambar 21. ERD ringkas (erd.mmd). Kolom lengkap ada di SQL.")

    add_heading(doc, "4.2 Detail tabel utama", 2)
    table(
        doc,
        ["Tabel", "Kolom penting", "Relasi"],
        [
            ["users", "email, password_hash, ktp_status, bank_status, wallet", "induk hampir semua"],
            ["categories", "name, parent_id, kind DIGITAL/FISIK", "self FK; ke services & jobs"],
            ["services", "seller_id, price, is_active, cover", "N jasa : 1 seller"],
            ["jobs", "buyer_id, budget, status OPEN/FILLED, is_active, is_urgent", "N lowongan : 1 klien"],
            ["applications", "job_id, seller_id, status, reject_kind, reject_reason", "unik user+job"],
            ["orders", "source SERVICE/JOB, status, escrow UNPAID/HELD/RELEASED", "buyer + seller"],
            ["payments", "order_id, status PENDING/PAID/EXPIRED", "1 PENDING & 1 PAID per order"],
            ["work_submissions", "file bukti, revision_no 1–3", "N : 1 order"],
            ["reviews", "rating 1–5, comment", "1 order, reviewer–reviewee"],
            ["user_portfolios", "title, file_url, category opsional", "maks 12 per user"],
        ],
    )

    add_heading(doc, "4.3 Alasan desain", 2)
    para(
        doc,
        "Escrow dipisah dari status pesanan supaya “sedang dikerjakan” tidak tercampur dengan "
        "“uang sudah ditahan”. Lowongan memisahkan status lamaran (OPEN/CLOSED/FILLED) dari "
        "is_active (tampil/sembunyi di katalog) — kalau disatukan, menutup deadline dan menyembunyikan "
        "kartu jadi bentrok. reject_kind membedakan tolak manual, otomatis karena orang lain diterima, "
        "dan otomatis karena order kedaluwarsa. Tidak pakai Sequelize: transaksi “terima pelamar + "
        "tolak sisanya + buat order” lebih jelas ditulis SQL langsung.",
    )

    add_heading(doc, "5. Deskripsi Teknis Mengenai Integrasi API", 1)
    add_heading(doc, "5.1 Spesifikasi API", 2)
    para(
        doc,
        "Base URL lokal http://localhost:3000. Marketplace di /api/mp, gateway di /gateway/api, "
        "webhook di /api/webhooks. Format JSON. Frontend Vite mem-proxy /api ke backend. "
        "Tidak ada XML.",
    )
    table(
        doc,
        ["Method", "Endpoint", "Kegunaan"],
        [
            ["POST", "/api/mp/auth/register", "daftar akun"],
            ["POST", "/api/mp/auth/login", "login, Set-Cookie token"],
            ["POST", "/api/mp/auth/logout", "hapus cookie"],
            ["GET", "/api/mp/me", "sesi user"],
            ["GET", "/api/mp/jasa", "katalog jasa (publik)"],
            ["POST", "/api/mp/jasa", "buat jasa (login + bank approved)"],
            ["GET", "/api/mp/lowongan", "katalog lowongan (publik)"],
            ["POST", "/api/mp/jasa/:id/sewa", "request sewa"],
            ["POST", "/api/mp/applications/:id/terima", "rekrut 1 pelamar"],
            ["POST", "/api/mp/orders/:id/bayar", "mulai pembayaran"],
            ["POST", "/gateway/api/transactions/:code/pay", "simulasi bayar"],
            ["POST", "/api/webhooks/payment-gateway", "konfirmasi PAID"],
        ],
    )

    add_heading(doc, "5.2 Autentikasi", 2)
    para(
        doc,
        "Token-based, tapi token tidak dikirim Bearer di header. JWT disimpan cookie httpOnly "
        "(name: token), dikirim otomatis same-origin. Payload berisi id user dan role. "
        "Admin bootstrap: admin@mail.com. Gateway merchant memakai API key di env, terpisah dari JWT user.",
    )
    shot(doc, "04-register", "Gambar 22. Form register — pintu masuk auth.")
    shot(doc, "02-login", "Gambar 23. Halaman login.")

    add_heading(doc, "5.3 Error handling", 2)
    para(
        doc,
        "Helper fail(res, status, message, errors) mengembalikan { ok:false, error, errors }. "
        "400 validasi form, 401 belum login / password salah, 403 salah orang atau akun banned, "
        "404 resource tidak ada. Frontend api.js melempar error ke Alert. Upload salah tipe/ukuran "
        "juga 400. Gateway gagal: order tetap UNPAID, user lihat pesan, bukan halaman kosong.",
    )
    shot(doc, "03-login-gagal", "Gambar 24. Login gagal — API 401, UI Alert “Email atau password salah”.")
    shot(doc, "05-register-validasi", "Gambar 25. Validasi frontend: konfirmasi password tidak cocok.")
    shot(doc, "23-order-tidak-ada", "Gambar 26. GET order id tidak ada — “Pesanan tidak ditemukan”.")

    add_heading(doc, "6. Hasil Pengujian dengan Berbagai Skenario", 1)
    add_heading(doc, "6.1 Metodologi", 2)
    para(
        doc,
        "Pengujian yang dikerjakan: (1) manual end-to-end di Chrome terhadap UI localhost:5173; "
        "(2) cek response API lewat browser/fetch karena auth memakai cookie; (3) cek constraint "
        "database (unik payment PAID). Unit test otomatis (Jest) belum ditulis — fokusnya alur "
        "yang rawan status dan uang. Screenshot di bawah diambil dari app yang sama, bukan mockup.",
    )

    add_heading(doc, "6.2 Skenario pengujian", 2)
    table(
        doc,
        ["Skenario", "Expected", "Actual"],
        [
            [
                "Login email/password salah",
                "Tidak masuk; pesan error",
                "401 + Alert “Email atau password salah”. Lihat Gbr. 16.",
            ],
            [
                "Register password ≠ konfirmasi",
                "Akun tidak dibuat",
                "Alert “Password dan konfirmasi password tidak cocok.” Gbr. 17.",
            ],
            [
                "Login admin@mail.com / admin123",
                "Masuk, cookie terpasang",
                "Redirect ke panel admin / dashboard. Gbr. 13 & 9.",
            ],
            [
                "Buka /jasa dan /lowongan (DB kosong)",
                "Empty state, bukan crash",
                "“Belum ada jasa/lowongan” + tombol post. Gbr. 7–8.",
            ],
            [
                "Post jasa judul kosong",
                "Browser/API menolak",
                "Required HTML “Please fill out this field”. Gbr. 19.",
            ],
            [
                "Akses /orders/99999",
                "404 ramah",
                "Kartu “Tidak Dapat Mengakses Pesanan”. Gbr. 18.",
            ],
            [
                "GET /api/mp/profile/1 (akun admin)",
                "Tidak dibuka sebagai storefront",
                "“Profil tidak ditemukan”. Gbr. 20.",
            ],
            [
                "Publish jasa tanpa rekening approved",
                "Guard seller",
                "Banner “Lengkapi rekening” di dashboard/form. Gbr. 9 & 11.",
            ],
            [
                "OTP email sudah verified",
                "Tidak minta OTP lagi",
                "Halaman “Email sudah terverifikasi.” Gbr. 21.",
            ],
            [
                "Terima 1 pelamar (alur kode)",
                "Pelamar lain AUTO_FILLED + order dibuat",
                "Diverifikasi di applicationFlow.js + constraint DB. Data listing di DB saat uji UI masih 0.",
            ],
            [
                "Bayar dua kali order sama",
                "Tidak dobel PAID",
                "Unik 1 PENDING & 1 PAID di tabel payments.",
            ],
        ],
    )

    add_heading(doc, "6.3 Transaksi berhasil dan gagal (screenshot)", 2)
    para(
        doc,
        "Lingkungan uji ini database-nya baru (katalog 0, pesanan 0). Jadi bukti UI yang bisa "
        "diambil ulang hari ini adalah transaksi auth dan guard: gagal login, gagal register, "
        "gagal buka order, gagal buka profil admin, form post ditolak browser. Alur escrow "
        "berhasil (bayar mock → HELD → setujui) sudah diuji sebelumnya di mesin yang sama; "
        "langkahnya ada di bagian 2.3. Untuk demo dosen, seed listing lalu ulangi sewa/bayar.",
    )
    shot(doc, "22-post-jasa-validasi", "Gambar 27. Transaksi gagal: publish jasa tanpa judul (validasi required).")
    shot(doc, "19-profil-publik", "Gambar 28. Transaksi gagal: profil publik id admin tidak diekspos.")
    shot(doc, "13-verifikasi-email", "Gambar 29. Transaksi berhasil: status email sudah terverifikasi.")
    shot(doc, "24-verifikasi-bank", "Gambar 30. Form rekening — syarat nama wajib sama dengan KTP.")

    add_heading(doc, "7. Panduan Penggunaan dan Instalasi", 1)
    add_heading(doc, "7.1 Persiapan lingkungan", 2)
    para(doc, "Perlu: Node.js 18+, MySQL/MariaDB (XAMPP boleh), browser Chrome.", justify=False)
    para(doc, "Database", bold=True, justify=False)
    para(
        doc,
        "Buat database proyek_marketplace. Import backend/database/phpmyadmin_full_override.sql "
        "secara utuh (DROP + CREATE). Jangan jalankan migrasi 001–014 setelah file itu. "
        "Admin seed: admin@mail.com / admin123.",
    )
    para(doc, "Konfigurasi backend/.env", bold=True, justify=False)
    para(
        doc,
        "DB_HOST, DB_USER, DB_PASSWORD, DB_NAME=proyek_marketplace, JWT_SECRET, "
        "MERCHANT_API_KEY, WEBHOOK_SECRET, GATEWAY_FRONTEND_URL=http://localhost:5173.",
    )

    add_heading(doc, "7.2 Menjalankan aplikasi", 2)
    para(
        doc,
        "Terminal 1: cd backend && npm install && npm run dev  → http://localhost:3000. "
        "Terminal 2: cd frontend && npm install && npm run dev  → http://localhost:5173 "
        "(proxy /api ke 3000). Buka beranda di browser, daftar atau login admin.",
    )

    add_heading(doc, "7.3 Akses API dan testing endpoint", 2)
    para(
        doc,
        "Karena JWT di cookie, tes paling gampang lewat browser (sudah login) atau Thunder Client "
        "dengan cookie. Bearer Authorization tidak dipakai untuk /api/mp. Contoh: POST /api/mp/auth/login "
        "body JSON { email, password } → Set-Cookie. Lalu GET /api/mp/jasa tanpa login. "
        "POST /api/mp/jasa butuh cookie + bank approved.",
    )

    add_heading(doc, "8. Kesimpulan dan Rekomendasi Pengembangan Lanjutan", 1)
    add_heading(doc, "8.1 Ringkasan pencapaian", 2)
    para(
        doc,
        "Yang jadi bukan cuma CRUD. Ada dua katalog, auth cookie, verifikasi identitas, dashboard, "
        "admin, chat, dan kerangka escrow. Manfaatnya: alur sewa/lamaran punya status yang bisa "
        "dijejak, bukan chat liar. Di lingkungan uji kosong pun UI tetap menampilkan empty state "
        "dan error yang jelas.",
    )
    add_heading(doc, "8.2 Pembelajaran teknis", 2)
    para(
        doc,
        "Yang kerasa: status order gampang kacau kalau dicampur dengan “tampil di katalog” "
        "(makanya jobs.is_active terpisah dari OPEN/CLOSED). Terima pelamar harus atomik. "
        "Cookie httpOnly lebih aman dari JWT di localStorage, tapi bikin tes Postman sedikit ribet. "
        "MySQL + mysql2 cukup untuk tugas ini tanpa ORM.",
    )
    add_heading(doc, "8.3 Rencana lanjutan", 2)
    bullet(
        doc,
        [
            "Deploy cloud + MySQL managed.",
            "Payment gateway asli (Midtrans/Xendit) menggantikan mock.",
            "Tes otomatis untuk escrow dan auto-reject pelamar.",
            "Email/SMS OTP sungguhan, bukan simulasi.",
        ],
    )

    add_heading(doc, "9. Lampiran", 1)
    add_heading(doc, "9.1 Screenshot tambahan", 2)
    shot(doc, "14-notifikasi", "Lampiran A. Halaman notifikasi (empty state).")
    shot(doc, "15-chat", "Lampiran B. Inbox chat (0 percakapan).")
    shot(doc, "17-admin-users", "Lampiran C. Admin — daftar pengguna marketplace.")
    shot(doc, "18-admin-orders", "Lampiran D. Admin — daftar pesanan.")
    shot(doc, "10-setelah-login", "Lampiran E. Setelah login admin: panel operasional.")

    add_heading(doc, "9.2 Kode / log pengujian", 2)
    para(
        doc,
        "Sumber diagram: docs/diagrams/*.mmd. Deskripsi langkah: docs/FLOW_SISTEM.md. "
        "Render diagram: DIAGRAMS_ONLY=1 node docs/capture_laporan.mjs. "
        "Builder laporan: docs/build_laporan.py. Contoh log API saat uji: GET /api/mp/jasa "
        "→ { ok:true, data:[] }. GET /api/mp/profile/1 → { ok:false, error:\"Profil tidak ditemukan\" }. "
        "Login salah → 401 Email atau password salah.",
    )
    para(
        doc,
        "Repo GitHub: https://github.com/spzhrrr/proyek_marketplace_jasa. "
        "Exam 1 adalah template slide; yang diikuti untuk dokumen ini Exam 3. "
        "Exam 2 (silabus) terpenuhi untuk Express+React+CRUD+auth+DB relasional; Sequelize/EJS/Redux/"
        "deploy cloud tidak 1:1 (diganti mysql2, React, Context, localhost).",
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("— selesai —")
    set_run_font(r, size=10, color=(148, 163, 184))

    OUT.parent.mkdir(parents=True, exist_ok=True)
    for path in (OUT, OUT_FALLBACK, OUT_FALLBACK2):
        try:
            doc.save(path)
            print(f"Wrote {path}")
            break
        except PermissionError:
            continue
    else:
        raise PermissionError("Semua file laporan Word terkunci. Tutup Word lalu jalankan ulang.")


if __name__ == "__main__":
    build()
